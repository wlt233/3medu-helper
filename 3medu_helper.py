#!/usr/bin/env python3
# coding=utf-8
import requests
import re, json, time, os, base64, math 
from docx import Document
#正则表达式匹配pattern
pattern1 = 'id="j_id1:javax.faces.ViewState:0" value="(.*?)" a'
pattern2 = 'id="j_id1:javax.faces.ClientWindow:0" value="(.*?)" a'
pattern3 = 'JSESSIONID=(.*?); P'
pattern4 = "'Location': '(.*?)', 'Content-Language'"
pattern5 = "csfcfc=(.*?); Path=/"

print("欢迎使用三米教育错题整理程序 v2.0 by某魏")

try:
    with open('setting.txt', 'r') as f:
        f.readline()
        stuid = f.readline().replace('\n','')
        f.readline()
        paswd = f.readline().replace('\n','')
except FileNotFoundError:
    with open('setting.txt', 'w+') as f:
        stuid = input("请输入您的id    ")
        paswd = input("请输入您的密码  ")
        f.write("#请在下一行键入你的id:\n")
        f.write(stuid+'\n')
        f.write("#请在下一行键入你的密码:\n")
        f.write(paswd+'\n')

print("正在使用 id = "+stuid+" 密码 = "+paswd+" 登陆")

session = requests.Session()
#cookies登陆 备用
'''
raw_cookies = "Hm_lvt_9cef549829657e6fb4e3fd3fe45166d1=1553782489; Hm_lpvt_9cef549829657e6fb4e3fd3fe45166d1=1553782489; JSESSIONID=91632eb0a80ddd8a4542ea202c19"
cookies={}
for line in raw_cookies.split(';'):
    key,value=line.split('=',1)
    cookies[key]=value
print(cookies)
'''


headers = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/x-www-form-urlencoded",
    "Host": "corework.3medu.com",
    "Origin": "https://corework.3medu.com"}
first_login = session.get("https://corework.3medu.com/login.html", headers=headers)
jfv = str(re.search(pattern1,first_login.text).group(1))
jfw = str(re.search(pattern2,first_login.text).group(1))
jid = str(re.search(pattern3,str(first_login.headers)).group(1))
#print(jfv,jfw,jid)


prelogin_data_str = "login-to-login=login-to-login&javax.faces.ViewState="+jfv.replace(':','%3A')+"&javax.faces.ClientWindow="+jfw.replace(':','%3A')+"&login-to-login%3Aj_idt16=login-to-login%3Aj_idt16"
prelogin = session.post("https://corework.3medu.com/login.html;jsessionid="+jid+"?jfwid="+jfw, data=prelogin_data_str, headers=headers)
jfv2 = str(re.search(pattern1,prelogin.text).group(1))
#print(jfv2)


'''
#这里的form data，用字典就各种500，str就没事...
login_data_dict = {'loginForm': 'loginForm', 
    'loginForm%3Aj_idt19': stuid, 
    'loginForm%3Aj_idt21': paswd, 
    'loginForm%3Aj_idt23': r'%E7%99%BB%E5%BD%95',
    'javax.faces.ViewState': jfv.replace(':','%3A'), 
    'javax.faces.ClientWindow': jfw.replace(':','%3A')
    }
'''
login_data_str = "loginForm=loginForm&loginForm%3Aj_idt19="+stuid+"&loginForm%3Aj_idt21="+paswd+r"&loginForm%3Aj_idt23=%E7%99%BB%E5%BD%95&javax.faces.ViewState="+jfv2.replace(':','%3A')+"&javax.faces.ClientWindow="+jfw.replace(':','%3A')
login = session.post("https://corework.3medu.com/f-login/f-login.html?jfwid="+jfw, data=login_data_str, headers=headers)
jfv3 = str(re.search(pattern1,login.text).group(1))
jfw2 = str(re.search(pattern3,str(login.headers)).group(1))+':0'
#print(jfv3,jfw2)


headers2 = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/x-www-form-urlencoded",
    "Host": "corework.3medu.com",
    "Origin": "https://corework.3medu.com",
    "Referer": "https://corework.3medu.com/f-login/f-login.html?jfwid="+jfw}
redirect_data_str = "login-to-redirect=login-to-redirect&javax.faces.ViewState="+jfv3.replace(':','%3A')+"&javax.faces.ClientWindow="+jfw.replace(':','%3A')+"&login-to-redirect%3Aj_idt9=login-to-redirect%3Aj_idt9"
redirect = session.post("https://corework.3medu.com/f-login/f-login-redirect.html?jfwid="+jfw, data=redirect_data_str, headers=headers2)#, cookies={"JSESSIONID": jfw2.replace(":0","")})
jfv4 = str(re.search(pattern1,redirect.text).group(1))





headers3 = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/x-www-form-urlencoded",
    "Host": "corework.3medu.com",
    "Origin": "https://corework.3medu.com",
    "Referer": "https://corework.3medu.com/portal/portal.html"+jfw,
    "Cookie": "JSESSIONID="+jfw2.replace(":0","")}
adls_data_str = "TO_ADLS=TO_ADLS&javax.faces.ViewState="+jfv4.replace(':','%3A')+"&javax.faces.ClientWindow="+jfw2.replace(":0",":2")+"&TO_ADLS%3Aj_idt16=TO_ADLS%3Aj_idt16"
redirect2 = session.post("https://corework.3medu.com/portal/portal.html?jfwid="+jfw2.replace(":0",":2"),  headers=headers3, data=adls_data_str, allow_redirects=False)
callbackurl = str(re.search(pattern4,str(redirect2.headers)).group(1))
redirect3 = session.get(callbackurl, headers=headers3, allow_redirects=False)
jfw3 = str(re.search(pattern3,str(redirect3.headers)).group(1))
#print(jfw3)



headers4 = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/x-www-form-urlencoded",
    "Host": "adls.3medu.com",
    "Referer": "https://corework.3medu.com/portal/portal.html",
    "cookie":"JSESSIONID="+jfw3}
redirect4 = session.get("https://adls.3medu.com/portal.html", headers=headers4, allow_redirects=False)
#print(redirect4.headers)
csfcfc1 = str(re.search(pattern5,str(redirect4.headers)).group(1))


adls_cookies1 = {'JSESSIONID':jfw3,'cscfcf':csfcfc1}
headers_adls = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/x-www-form-urlencoded",
    "Referer": "https://corework.3medu.com/portal/portal.html",
    "Host": "adls.3medu.com"}
portal = session.get("https://adls.3medu.com/portal.html?r=1", headers=headers_adls, cookies=adls_cookies1)
csfcfc = str(re.search(pattern5,str(portal.headers)).group(1))


adls_cookies = {'JSESSIONID':jfw3,'csfcfc':csfcfc}
headers_adls = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "application/json, text/javascript, */*; q=0.01",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/json;charset=UTF-8",
    "Origin": "https://adls.3medu.com",
    "X-Requested-With": "XMLHttpRequest",
    "Host": "adls.3medu.com"}


headers_edbes = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "application/json, text/javascript, */*; q=0.01",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/json;charset=UTF-8",
    "Origin": "https://adls.3medu.com",
    "Referer": "https://corework.3medu.com/portal/portal.html",
    "X-Requested-With": "XMLHttpRequest",
    "Host": "adls.3medu.com"}
edbes_data = {"status":"ALL"}
edbes = session.post("https://adls.3medu.com/rs/exam/extractDatesByExamStatus", headers=headers_edbes, cookies=adls_cookies, json=edbes_data)
#print(edbes.text)



headers_febd = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36",
    "Upgrade-Insecure-Requests":"1",
    "Accept": "application/json, text/javascript, */*; q=0.01",
    "Accept-Encoding":"gzip, deflate, br",
    "Accept-Language": "zh-CN,zh;q=0.9",
    "Cache-Control": "max-age=0",
    "Connection": "keep-alive",
    "Content-Type": "application/json;charset=UTF-8",
    "Origin": "https://adls.3medu.com",
    "Referer": "https://adls.3medu.com/portal.html?r=1",
    "X-Requested-With": "XMLHttpRequest",
    "Host": "adls.3medu.com"}
febd_data = {"first":0,"amount":100,"status":"ALL","orderType":"DATE","ascending":"false"}
febd = session.post("https://adls.3medu.com/rs/exam/findExamsByDate", headers=headers_febd, cookies=adls_cookies, json=febd_data)
#print(febd.text)
examlist = json.loads(febd.text)["content"]



info = session.get("https://adls.3medu.com/rs/user/findBasicInfo",headers=headers_febd, cookies=adls_cookies)
userinfo = json.loads(info.text)["content"]
username = userinfo["lastName"]+userinfo["firstName"]
userid = userinfo["identityCode"]
print("欢迎您！  ",username,userid)
print(' ')



print("查询到以下测试：")
exami = 1
for eachexam in examlist:
    print('%2d'%(exami),eachexam['paperName'])
    exami += 1
print(' ')
f = 0
while f==0:
    chosen = input("请选择测试，键入测试名前的数字并回车确定：")
    chosenname=examlist[int(chosen)-1]['paperName']
    chosenid=examlist[int(chosen)-1]['examineeId']
    print(' ')
    print("您选择的是：",chosenname)
    chosenexam = session.get("https://adls.3medu.com/rs/exam/getWrongItems?id="+str(chosenid), headers=headers_adls, cookies=adls_cookies)
    try:
        examcontent = json.loads(chosenexam.text)["content"]
        f = 1
    except (NameError,KeyError):
        print("您还未完成这次测试！")
        print(' ')

print("您做错了以下题目：")
wronglist = []
for eachques in examcontent:
    solution = session.get("https://adls.3medu.com/rs/exam/getSolution?id="+str(eachques["itemId"]), headers=headers_adls, cookies=adls_cookies)
    solution_dict = json.loads(solution.text)["content"]
    optionslist = solution_dict["options"]
    for eachop in range(0,len(optionslist)):
        if optionslist[eachop]["rightChoice"] == False and optionslist[eachop]["selected"] == True:
            print(solution_dict['code'], optionslist[eachop]["seqCode"])
            wronglist.append(solution_dict)


document = Document()
document.add_heading(chosenname, 0) 
document.add_heading(str(username)+'  '+str(userid), 2)
if not os.path.exists('./latex'):
    os.makedirs('./latex')

session.close()
session2 = requests.Session()
headerslatex = {"User-Agent":"Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/69.0.3497.100 Safari/537.36"}

for eachsolution in wronglist:
    quesid = eachsolution['code']
    latexlist,imglist = [],[]
    print("正在处理 "+quesid+" 请稍等！")

    latexi = 1
    for eachqueslatex in eachsolution['content']['latexs']:
        lateximg = session2.get("http://latex.codecogs.com/gif.latex?"+eachqueslatex, headers=headerslatex)
        with open('./latex/'+str(quesid)+'_ques_'+str(latexi)+'.gif', 'wb+') as f:
            f.write(lateximg.content)
            f.close()
            latexlist.append('./latex/'+str(quesid)+'_ques_'+str(latexi)+'.gif')
            print(str(quesid)+'_ques_'+str(latexi)+'.gif saved.')
            latexi += 1
    imgi = 1
    for eachquesimg in eachsolution['content']['images']:
        with open('./latex/'+str(quesid)+'_quesimg_'+str(imgi)+'.png','wb+') as f:
            f.write(base64.b64decode(eachquesimg['src'].replace('data:image/png;base64,','')))
            f.close()
            imglist.append('./latex/'+str(quesid)+'_quesimg_'+str(imgi)+'.png')
            print(str(quesid)+'_quesimg_'+str(imgi)+'.png saved.')
    questextraw = eachsolution['content']['content'].replace('{$$<I>$$}','')
    questext = []
    questexti = 0
    questextj = 0
    for i in range(0,len(latexlist)):
        questexti = questextraw.find('{$$<L>$$}',questextj)
        questext.append(questextraw[questextj:questexti])
        questext.append("{$$<L>$$}")
        questextj = questexti+9
    questext.append(questextraw[questextj:len(questextraw)])
    #print(questext)
    doclatexi = 0
    questextdoc = document.add_paragraph('')
    eachquesrun = questextdoc.add_run(quesid)
    eachquesrun.font.size = 180000
    eachquesrun.bold,eachquesrun.italic = True,True
    for eachquestext in questext:
        if eachquestext != '{$$<L>$$}':
            eachquesrun = questextdoc.add_run(eachquestext)
        else:
            pic = eachquesrun.add_picture(latexlist[doclatexi])
            picheight,picwidth = pic.height,pic.width
            pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)
            questextdoc.add_run('   ')
            doclatexi += 1
    for eachdocqimg in imglist:
        pic = document.add_picture(eachdocqimg)
        picheight,picwidth = pic.height,pic.width
        pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)
    #1px=12700

    for eachoption in eachsolution['options']:
        seq = eachoption['seqCode']
        oplatexlist,opimglist = [],[]
        latexi = 1
        for eachopslatex in eachoption['content']['latexs']:
            lateximg = session2.get("http://latex.codecogs.com/gif.latex?"+eachopslatex, headers=headerslatex)
            with open('./latex/'+str(quesid)+'_op_'+seq+str(latexi)+'.gif', 'wb+') as f:
                f.write(lateximg.content)
                f.close()
                oplatexlist.append('./latex/'+str(quesid)+'_op_'+seq+str(latexi)+'.gif')
                print(str(quesid)+'_op_'+seq+str(latexi)+'.gif saved.')
                latexi += 1
        imgi = 1
        for eachquesimg in eachoption['content']['images']:
            with open('./latex/'+str(quesid)+'_opimg_'+seq+str(imgi)+'.png','wb+') as f:
                f.write(base64.b64decode(eachquesimg['src'].replace('data:image/png;base64,','')))
                f.close()
                opimglist.append('./latex/'+str(quesid)+'_opimg_'+str(imgi)+'.png')
                print(str(quesid)+'_opimg_'+str(imgi)+'.png saved.')
        optextraw = eachoption['content']['content'].replace('{$$<I>$$}','')
        optext = []
        optexti = 0
        optextj = 0
        for i in range(0,len(oplatexlist)):
            optexti = optextraw.find('{$$<L>$$}',optextj)
            optext.append(optextraw[optextj:optexti])
            optext.append("{$$<L>$$}")
            optextj = optexti+9
        optext.append(optextraw[optextj:len(optextraw)])
        docoplatexi = 0
        optextdoc = document.add_paragraph('')
        eachoprun = optextdoc.add_run(seq+". ")
        eachoprun.font.size = 160000
        eachoprun.bold = True
        for eachoptext in optext:
            if eachoptext != '{$$<L>$$}':
                eachoprun = optextdoc.add_run(eachoptext)
            else:
                pic = eachoprun.add_picture(oplatexlist[docoplatexi])
                picheight,picwidth = pic.height,pic.width
                pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)
                optextdoc.add_run('   ')
                docoplatexi += 1
        for eachdocqimg in opimglist:
            pic = document.add_picture(eachdocqimg)
            picheight,picwidth = pic.height,pic.width
            pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)
        if eachoption['rightChoice'] == True:
            eachoprun = optextdoc.add_run("(正确答案)")
        if eachoption['selected'] == True:
            eachoprun = optextdoc.add_run("(错误选择)")


    for eachsolu in eachsolution['solutions']:
        solulatexlist,soluimglist = [],[]
        latexi = 1
        for eachopslatex in eachsolu['latexs']:
            lateximg = session2.get("http://latex.codecogs.com/gif.latex?"+eachopslatex, headers=headerslatex)
            with open('./latex/'+str(quesid)+'_solu_'+seq+str(latexi)+'.gif', 'wb+') as f:
                f.write(lateximg.content)
                f.close()
                solulatexlist.append('./latex/'+str(quesid)+'_solu_'+seq+str(latexi)+'.gif')
                print(str(quesid)+'_solu_'+seq+str(latexi)+'.gif saved.')
                latexi += 1
        imgi = 1
        for eachquesimg in eachsolu['images']:
            with open('./latex/'+str(quesid)+'_soluimg_'+seq+str(imgi)+'.png','wb+') as f:
                f.write(base64.b64decode(eachquesimg['src'].replace('data:image/png;base64,','')))
                f.close()
                soluimglist.append('./latex/'+str(quesid)+'_soluimg_'+str(imgi)+'.png')
                print(str(quesid)+'_soluimg_'+str(imgi)+'.png saved.')
        solutextraw = eachsolu['content'].replace('{$$<I>$$}','')
        solutext = []
        solutexti = 0
        solutextj = 0
        for i in range(0,len(solulatexlist)):
            solutexti = solutextraw.find('{$$<L>$$}',solutextj)
            solutext.append(solutextraw[solutextj:solutexti])
            solutext.append("{$$<L>$$}")
            solutextj = solutexti+9
        solutext.append(solutextraw[solutextj:len(solutextraw)])
        docsolulatexi = 0
        solutextdoc = document.add_paragraph('')
        eachsolurun = solutextdoc.add_run("解析:  ")
        eachsolurun.font.size = 160000
        eachsolurun.bold = True
        for eachsolutext in solutext:
            if eachsolutext != '{$$<L>$$}':
                eachsolurun = solutextdoc.add_run(eachsolutext)
            else:
                pic = eachsolurun.add_picture(solulatexlist[docsolulatexi])
                picheight,picwidth = pic.height,pic.width
                pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)
                solutextdoc.add_run('   ')
                docsolulatexi += 1
        for eachdocqimg in soluimglist:
            pic = document.add_picture(eachdocqimg)
            picheight,picwidth = pic.height,pic.width
            pic.height,pic.width=math.floor(picheight/1.5),math.floor(picwidth/1.5)

print(chosenname+".docx已保存  错题整理完成")
document.save(chosenname+'.docx')  
